#本程式為因應環保署專案，需要產出全國所有測站的基本資料表
#先利用python-docx設計好一個WORD表格，並且將事先匯整好的資料寫入表格中
#以迴圈批次處理，將這些資料匯整至一個(.docx)檔案內
#註：檔案中所使用到的資料皆為公開資料

from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement #設定網底
from docx.enum.text import WD_ALIGN_PARAGRAPH #水平置中
from docx.enum.table import WD_ALIGN_VERTICAL #垂直置中
import pandas as pd

document = Document() #開啟空白檔案
document.sections[0].page_height = Cm(29.7)  # 設定A4紙的高度
document.sections[0].page_width = Cm(21)  # 設定A4紙的寬
document.save('環保署測站基本資料表_年報用.docx') #儲存空白檔案
#########################讀取基本資料檔案#############################
station_info = pd.read_csv('環保署測站基本資料_填表用.csv')
#########################將基本資料各項目逐一列出#############################
station0 = list(station_info.iloc[:, 0])

for i in station0:
    station = station_info[(station_info.測站名稱 == str(i))]
    station_name = station.iloc[:, 0] #測站名稱
    station_site = station.iloc[:, 1] #測站位置
    aqi_area = station.iloc[:, 2] #空品區
    start_time = station.iloc[:, 3] #運作日期開始
    main_class = station.iloc[:, 4] #測站類別(主)
    secondary_class = station.iloc[:, 5] #測站類別(輔)
    county = station.iloc[:, 6] #縣市
    district = station.iloc[:, 7] #鄉鎮區
    address = station.iloc[:, 8] #地址
    lat = station.iloc[:, 9] #北緯
    lon = station.iloc[:, 10] #東經

    item_1_brand = station.iloc[:, 11] #CO儀器廠牌
    item_2_brand = station.iloc[:, 12] #NO2儀器廠牌
    item_3_brand = station.iloc[:, 13] #O3儀器廠牌
    item_4_brand = station.iloc[:, 14] #PM10儀器廠牌
    item_5_brand = station.iloc[:, 15] #PM2.5儀器廠牌
    item_6_brand = station.iloc[:, 16] #SO2儀器廠牌
    item_7_brand = station.iloc[:, 17] #NMHC儀器廠牌
    atm_item_1_brand = station.iloc[:, 18] #AMB_TEMP儀器廠牌
    atm_item_2_brand = station.iloc[:, 19] #RAINFALL儀器廠牌
    atm_item_3_brand = station.iloc[:, 20] #RH儀器廠牌
    atm_item_4_brand = station.iloc[:, 21] #WD_HR儀器廠牌
    atm_item_5_brand = station.iloc[:, 22] #WS_HR儀器廠牌
    atm_item_6_brand = station.iloc[:, 23] #WIND_DIREC儀器廠牌
    atm_item_7_brand = station.iloc[:, 24] #WIND_SPEED儀器廠牌
    atm_item_8_brand = station.iloc[:, 25] #RAIN_COND儀器廠牌
    atm_item_9_brand = station.iloc[:, 26] #PH_RAIN儀器廠牌
    atm_item_10_brand = station.iloc[:, 27] #UVB儀器廠牌
    other_item = station.iloc[:, 28] #其它監測項目
    pm25_manual = station.iloc[:, 29] #細懸浮微粒手動監測
    altitude = station.iloc[:, 30] #測站高度
    sample_connection_height = station.iloc[:, 31] #採樣口高度(m)
    sample_connection_angle = station.iloc[:, 32] #採樣口氣流角度 
    road_name1 = station.iloc[:, 33] #國道、省道
    road_distance1 = station.iloc[:, 34] #國道、省道距離
    road_name2 = station.iloc[:, 35] #一般道路
    road_distance2 = station.iloc[:, 36] #一般道路距離

    town_population = station.iloc[:, 37] #所在鄉鎮市區人口數量(人)
    town_population_density = station.iloc[:, 38] #所在鄉鎮市區人口密度(人/平方公里)

    document = Document('環保署測站基本資料表_年報用.docx') #開啟檔案

    table_1 = document.add_table(rows=30, cols=5) #設定初始表格大小

    c00 = table_1.cell(0, 0)
    c01 = table_1.cell(0, 1)
    c02 = table_1.cell(0, 2)
    c03 = table_1.cell(0, 3)
    c04 = table_1.cell(0, 4)
    c0 = c00.merge(c01).merge(c02).merge(c03).merge(c04)
    c0.text = '基本資訊'
    c0_xml_element = c0._tc #設定網底
    c0_properties = c0_xml_element.get_or_add_tcPr()
    shade_object = OxmlElement('w:shd')
    shade_object.set(qn('w:fill'), '#BDD6EE')
    c0_properties.append(shade_object)

    c10 = table_1.cell(1, 0)
    c11 = table_1.cell(1, 1)
    c12 = table_1.cell(1, 2)
    c13 = table_1.cell(1, 3)
    c14 = table_1.cell(1, 4)
    c1 = c11.merge(c12).merge(c13).merge(c14)
    c10.text = '測站名稱'
    c1.text = station_name.astype(str)

    c20 = table_1.cell(2, 0)
    c21 = table_1.cell(2, 1)
    c22 = table_1.cell(2, 2)
    c23 = table_1.cell(2, 3)
    c24 = table_1.cell(2, 4)
    c2 = c21.merge(c22).merge(c23).merge(c24)
    c20.text = '測站位置'
    c2.text = station_site.astype(str)

    c30 = table_1.cell(3, 0)
    c31 = table_1.cell(3, 1)
    c32 = table_1.cell(3, 2)
    c33 = table_1.cell(3, 3)
    c34 = table_1.cell(3, 4)
    c3 = c33.merge(c34)
    c30.text = '空品區'
    c32.text = '運作日期開始'
    c31.text = aqi_area.astype(str)
    c3.text = start_time.astype(str)

    c40 = table_1.cell(4, 0)
    c41 = table_1.cell(4, 1)
    c42 = table_1.cell(4, 2)
    c43 = table_1.cell(4, 3)
    c44 = table_1.cell(4, 4)
    c4 = c43.merge(c44)
    c40.text = '測站類別(主)'
    c42.text = '測站類別(輔)'
    c41.text = main_class.astype(str)
    c4.text = secondary_class.astype(str)

    c50 = table_1.cell(5, 0)
    c51 = table_1.cell(5, 1)
    c52 = table_1.cell(5, 2)
    c53 = table_1.cell(5, 3)
    c54 = table_1.cell(5, 4)
    c5 = c53.merge(c54)
    c50.text = '縣市'
    c52.text = '鄉鎮區'
    c51.text = county.astype(str)
    c5.text = district.astype(str)

    c60 = table_1.cell(6, 0)
    c61 = table_1.cell(6, 1)
    c62 = table_1.cell(6, 2)
    c63 = table_1.cell(6, 3)
    c64 = table_1.cell(6, 4)
    c6 = c61.merge(c62).merge(c63).merge(c64)
    c60.text = '地址'
    c6.text = address.astype(str)

    c70 = table_1.cell(7, 0)
    c71 = table_1.cell(7, 1)
    c72 = table_1.cell(7, 2)
    c73 = table_1.cell(7, 3)
    c74 = table_1.cell(7, 4)
    c7 = c71.merge(c72).merge(c73).merge(c74)
    c70.text = '經緯度'
    c7.text = lat.astype(str)

    c80 = table_1.cell(8, 0)
    c81 = table_1.cell(8, 1)
    c82 = table_1.cell(8, 2)
    c83 = table_1.cell(8, 3)
    c84 = table_1.cell(8, 4)
    c8 = c80.merge(c81).merge(c82).merge(c83).merge(c84)
    c8.text = '空品監測項目及儀器廠牌'
    c8_xml_element = c8._tc #設定網底
    c8_properties = c8_xml_element.get_or_add_tcPr()
    shade_object = OxmlElement('w:shd')
    shade_object.set(qn('w:fill'), '#FFE599')
    c8_properties.append(shade_object)

    c90 = table_1.cell(9, 0)
    c91 = table_1.cell(9, 1)
    c92 = table_1.cell(9, 2)
    c93 = table_1.cell(9, 3)
    c94 = table_1.cell(9, 4)
    c9 = c93.merge(c94)
    c90.text = 'CO'
    c92.text = 'NO'
    c92_subscript = c92.paragraphs[0].add_run('2').font.subscript = True
    c91.text = item_1_brand.astype(str)
    c9.text = item_2_brand.astype(str)

    c10_0 = table_1.cell(10, 0)
    c10_1 = table_1.cell(10, 1)
    c10_2 = table_1.cell(10, 2)
    c10_3 = table_1.cell(10, 3)
    c10_4 = table_1.cell(10, 4)
    c10 = c10_3.merge(c10_4)
    c10_0.text = 'O'
    c10_2.text = 'PM'
    c10_0_subscript = c10_0.paragraphs[0].add_run('3').font.subscript = True
    c10_2_subscript = c10_2.paragraphs[0].add_run('10').font.subscript = True
    c10_1.text = item_3_brand.astype(str)
    c10.text = item_4_brand.astype(str)

    c11_1 = table_1.cell(11, 1)
    c11_2 = table_1.cell(11, 2)
    c11_3 = table_1.cell(11, 3)
    c11_4 = table_1.cell(11, 4)
    c11 = c11_3.merge(c11_4)
    rows11_cells = table_1.rows[11].cells
    rows11_cells[0].text = 'PM'
    rows11_cells[2].text = 'SO'
    c11_0_subscript = rows11_cells[0].paragraphs[0].add_run('2.5').font.subscript = True
    c11_2_subscript = rows11_cells[2].paragraphs[0].add_run('2').font.subscript = True
    c11_1.text = item_5_brand.astype(str)
    c11.text = item_6_brand.astype(str)

    c12_1 = table_1.cell(12, 1)
    c12_2 = table_1.cell(12, 2)
    c12_3 = table_1.cell(12, 3)
    c12_4 = table_1.cell(12, 4)
    c12 = c12_3.merge(c12_4)
    rows12_cells = table_1.rows[12].cells
    rows12_cells[0].text = 'NMHC'
    c12_1.text = item_7_brand.astype(str)

    c13_0 = table_1.cell(13, 0)
    c13_1 = table_1.cell(13, 1)
    c13_2 = table_1.cell(13, 2)
    c13_3 = table_1.cell(13, 3)
    c13_4 = table_1.cell(13, 4)
    c13 = c13_0.merge(c13_1).merge(c13_2).merge(c13_3).merge(c13_4)
    c13.text = '氣象監測項目及儀器廠牌'
    c13_xml_element = c13._tc #設定網底
    c13_properties = c13_xml_element.get_or_add_tcPr()
    shade_object = OxmlElement('w:shd')
    shade_object.set(qn('w:fill'), '#FFE599')
    c13_properties.append(shade_object)

    c14_1 = table_1.cell(14, 1)
    c14_2 = table_1.cell(14, 2)
    c14_3 = table_1.cell(14, 3)
    c14_4 = table_1.cell(14, 4)
    c14 = c14_3.merge(c14_4)
    rows14_cells = table_1.rows[14].cells
    rows14_cells[0].text = 'AMB_TEMP'
    rows14_cells[2].text = 'RAINFALL'
    c14_1.text = atm_item_1_brand.astype(str)
    c14.text = atm_item_2_brand.astype(str)


    c15_1 = table_1.cell(15, 1)
    c15_2 = table_1.cell(15, 2)
    c15_3 = table_1.cell(15, 3)
    c15_4 = table_1.cell(15, 4)
    c15 = c15_3.merge(c15_4)
    rows15_cells = table_1.rows[15].cells
    rows15_cells[0].text = 'RH'
    rows15_cells[2].text = 'WD_HR'
    c15_1.text = atm_item_3_brand.astype(str)
    c15.text = atm_item_4_brand.astype(str)

    c16_1 = table_1.cell(16, 1)
    c16_2 = table_1.cell(16, 2)
    c16_3 = table_1.cell(16, 3)
    c16_4 = table_1.cell(16, 4)
    c16 = c16_3.merge(c16_4)
    rows16_cells = table_1.rows[16].cells
    rows16_cells[0].text = 'WS_HR'
    rows16_cells[2].text = 'WIND_DIREC'
    c16_1.text = atm_item_5_brand.astype(str)
    c16.text = atm_item_6_brand.astype(str)

    c17_1 = table_1.cell(17, 1)
    c17_2 = table_1.cell(17, 2)
    c17_3 = table_1.cell(17, 3)
    c17_4 = table_1.cell(17, 4)
    c17 = c17_3.merge(c17_4)
    rows17_cells = table_1.rows[17].cells
    rows17_cells[0].text = 'WIND_SPEED'
    rows17_cells[2].text = 'RAIN_COND'
    c17_1.text = atm_item_7_brand.astype(str)
    c17.text = atm_item_8_brand.astype(str)

    c18_1 = table_1.cell(18, 1)
    c18_2 = table_1.cell(18, 2)
    c18_3 = table_1.cell(18, 3)
    c18_4 = table_1.cell(18, 4)
    c18 = c18_3.merge(c18_4)
    rows18_cells = table_1.rows[18].cells
    rows18_cells[0].text = 'PH_RAIN'
    rows18_cells[2].text = 'UVB'
    c18_1.text = atm_item_9_brand.astype(str)
    c18.text = atm_item_10_brand.astype(str)

    c19_0 = table_1.cell(19, 0)
    c19_1 = table_1.cell(19, 1)
    c19_2 = table_1.cell(19, 2)
    c19_3 = table_1.cell(19, 3)
    c19_4 = table_1.cell(19, 4)
    c19 = c19_0.merge(c19_1).merge(c19_2).merge(c19_3).merge(c19_4)
    c19.text = '其他項目'
    c19_xml_element = c19._tc #設定網底
    c19_properties = c19_xml_element.get_or_add_tcPr()
    shade_object = OxmlElement('w:shd')
    shade_object.set(qn('w:fill'), '#FFE599')
    c19_properties.append(shade_object)

    c20_0 = table_1.cell(20, 0)
    c20_1 = table_1.cell(20, 1)
    c20_2 = table_1.cell(20, 2)
    c20_3 = table_1.cell(20, 3)
    c20_4 = table_1.cell(20, 4)   
    c20_merge0 = c20_0.merge(c20_1)
    c20_merge1 = c20_2.merge(c20_3).merge(c20_4)
    c20_merge0.text = '其他監測項目'
    c20_merge1.text = other_item.astype(str)

    c21_0 = table_1.cell(21, 0)
    c21_1 = table_1.cell(21, 1)
    c21_2 = table_1.cell(21, 2)
    c21_3 = table_1.cell(21, 3)
    c21_4 = table_1.cell(21, 4)   
    c21_merge0 = c21_0.merge(c21_1)
    c21_merge1 = c21_2.merge(c21_3).merge(c21_4)
    c21_merge0.text = '細懸浮微粒手動監測'
    c21_merge1.text = pm25_manual.astype(str)

    c22_0 = table_1.cell(22, 0)
    c22_1 = table_1.cell(22, 1)
    c22_2 = table_1.cell(22, 2)
    c22_3 = table_1.cell(22, 3)
    c22_4 = table_1.cell(22, 4)
    c22 = c22_0.merge(c22_1).merge(c22_2).merge(c22_3).merge(c22_4)
    c22.text = '測站周遭環境'
    c22_xml_element = c22._tc #設定網底
    c22_properties = c22_xml_element.get_or_add_tcPr()
    shade_object = OxmlElement('w:shd')
    shade_object.set(qn('w:fill'), '#A8D08D')
    c22_properties.append(shade_object)

    c23_0 = table_1.cell(23, 0)
    c23_1 = table_1.cell(23, 1)
    c23_2 = table_1.cell(23, 2)
    c23_3 = table_1.cell(23, 3)
    c23_4 = table_1.cell(23, 4)
    c23 = c23_3.merge(c23_4)
    c23_0.text = '測站高度(m)'
    c23_2.text = '採樣口高度(m)'
    c23_1.text = altitude.astype(str)
    c23.text = sample_connection_height.astype(str)

    c24_0 = table_1.cell(24, 0)
    c24_1 = table_1.cell(24, 1)
    c24_2 = table_1.cell(24, 2)
    c24_3 = table_1.cell(24, 3)
    c24_4 = table_1.cell(24, 4)   
    c24 = c24_3.merge(c24_4)
    c24_0.text = '採樣口氣流角度(°)'
    c24_1.text = sample_connection_angle.astype(str)

    rows25_cells = table_1.rows[25].cells
    rows25_cells[1].text = '國道、省道'
    rows25_cells[3].text = '距離(km)'
    rows25_cells[2].text = road_name1.astype(str)
    rows25_cells[4].text = road_distance1.astype(str)

    rows26_cells = table_1.rows[26].cells
    rows26_cells[1].text = '一般道路'
    rows26_cells[3].text = '距離(m)'
    rows26_cells[2].text = road_name2.astype(str)
    rows26_cells[4].text = road_distance2.astype(str)

    c25_0 = table_1.cell(25, 0)
    c26_0 = table_1.cell(26, 0)
    c25_26 = c25_0.merge(c26_0)
    c25_0.text = '最近主要道路'
    c25_0.paragraphs[0].add_run('').add_break()
    c25_0.paragraphs[0].add_run('及其距離').alignment = WD_ALIGN_PARAGRAPH.CENTER

    c27_0 = table_1.cell(27, 0)
    c27_1 = table_1.cell(27, 1)
    c27_2 = table_1.cell(27, 2)
    c27_3 = table_1.cell(27, 3)
    c27_4 = table_1.cell(27, 4)
    c27 = c27_0.merge(c27_1).merge(c27_2).merge(c27_3).merge(c27_4)
    c27.text = '所在區域人口資訊'
    c27_xml_element = c27._tc #設定網底
    c27_properties = c27_xml_element.get_or_add_tcPr()
    shade_object = OxmlElement('w:shd')
    shade_object.set(qn('w:fill'), '#F7CAAC')
    c27_properties.append(shade_object)

    c28_0 = table_1.cell(28, 0)
    c28_1 = table_1.cell(28, 1)
    c28_2 = table_1.cell(28, 2)
    c28_3 = table_1.cell(28, 3)
    c28_4 = table_1.cell(28, 4)
    c28 = c28_1.merge(c28_2).merge(c28_3).merge(c28_4)
    c28_0.text = '所在鄉鎮市區'
    c28_0.paragraphs[0].add_run('').add_break()
    c28_0.paragraphs[0].add_run('人口數量(人)').alignment = WD_ALIGN_PARAGRAPH.CENTER
    c28.text = town_population.astype(str)

    c29_0 = table_1.cell(29, 0)
    c29_1 = table_1.cell(29, 1)
    c29_2 = table_1.cell(29, 2)
    c29_3 = table_1.cell(29, 3)
    c29_4 = table_1.cell(29, 4)
    c29 = c29_1.merge(c29_2).merge(c29_3).merge(c29_4)
    c29_0.text = '所在鄉鎮市區人口'
    c29_0.paragraphs[0].add_run('').add_break()
    c29_0.paragraphs[0].add_run('密度(人/平方公里)').alignment = WD_ALIGN_PARAGRAPH.CENTER
    c29.text = town_population_density.astype(str)

    table_1.cell(26, 0).width = Inches(2.5)
    table_1.cell(26, 1).width = Inches(2.5)
    table_1.cell(26, 2).width = Inches(2.5)
    table_1.cell(26, 3).width = Inches(1.25)
    table_1.cell(26, 4).width = Inches(1.25)

    document.add_page_break()

    table_2 = document.add_table(rows=22, cols=7) #設定初始表格大小

    c00 = table_2.cell(0, 0)
    c01 = table_2.cell(0, 1)
    c02 = table_2.cell(0, 2)
    c03 = table_2.cell(0, 3)
    c04 = table_2.cell(0, 4)
    c05 = table_2.cell(0, 5)
    c06 = table_2.cell(0, 6)
    c0 = c00.merge(c01).merge(c02).merge(c03).merge(c04).merge(c05).merge(c06)
    c0.text = '鄰近10公里排放資訊'
    c0_xml_element = c0._tc #設定網底
    c0_properties = c0_xml_element.get_or_add_tcPr()
    shade_object = OxmlElement('w:shd')
    shade_object.set(qn('w:fill'), '#B4C6E7')
    c0_properties.append(shade_object)

    c10 = table_2.cell(1, 0)
    c11 = table_2.cell(1, 1)
    c12 = table_2.cell(1, 2)
    c13 = table_2.cell(1, 3)
    c14 = table_2.cell(1, 4)
    c15 = table_2.cell(1, 5)
    c16 = table_2.cell(1, 6)
    c11_12 = c11.merge(c12)
    c10.text = '各類型排放總量'
    c11_12.text = '物種'
    c13.text = 'SOx'
    c14.text = 'NOx'
    c15.text = 'THC'
    c16.text = 'TSP'

    c20 = table_2.cell(2, 0)
    c21 = table_2.cell(2, 1)
    c22 = table_2.cell(2, 2)
    c21_22 = c21.merge(c22)
    c21_22.text = '點源(T/y)'

    c30 = table_2.cell(3, 0)
    c31 = table_2.cell(3, 1)
    c32 = table_2.cell(3, 2)
    c31_32 = c31.merge(c32)
    c31_32.text = '線源(T/y)'

    c40 = table_2.cell(4, 0)
    c41 = table_2.cell(4, 1)
    c42 = table_2.cell(4, 2)
    c41_42 = c41.merge(c42)
    c41_42.text = '面源(T/y)'

    c10_20_30_40 = c10.merge(c20).merge(c30).merge(c40)

    table_2.cell(2, 3).text = station.iloc[:, 39].astype(str) #P_SOX總量
    table_2.cell(2, 4).text = station.iloc[:, 40].astype(str) #P_NOX總量
    table_2.cell(2, 5).text = station.iloc[:, 41].astype(str) #P_THC總量
    table_2.cell(2, 6).text = station.iloc[:, 42].astype(str) #P_TSP總量

    table_2.cell(3, 3).text = station.iloc[:, 43].astype(str) #L_SOX總量
    table_2.cell(3, 4).text = station.iloc[:, 44].astype(str) #L_NOX總量
    table_2.cell(3, 5).text = station.iloc[:, 45].astype(str) #L_THC總量
    table_2.cell(3, 6).text = station.iloc[:, 46].astype(str) #L_TSP總量

    table_2.cell(4, 3).text = station.iloc[:, 47].astype(str) #A_SOX總量
    table_2.cell(4, 4).text = station.iloc[:, 48].astype(str) #A_NOX總量
    table_2.cell(4, 5).text = station.iloc[:, 49].astype(str) #A_THC總量
    table_2.cell(4, 6).text = station.iloc[:, 50].astype(str) #A_TSP總量

    c50 = table_2.cell(5, 0)
    c51 = table_2.cell(5, 1)
    c52 = table_2.cell(5, 2)
    c53 = table_2.cell(5, 3)
    c54 = table_2.cell(5, 4)
    c55 = table_2.cell(5, 5)
    c56 = table_2.cell(5, 6)
    c51_52 = c51.merge(c52)
    c50.text = '各類型排放佔比'
    c51_52.text = '物種'
    c53.text = 'SOx'
    c54.text = 'NOx'
    c55.text = 'THC'
    c56.text = 'TSP'

    c60 = table_2.cell(6, 0)
    c61 = table_2.cell(6, 1)
    c62 = table_2.cell(6, 2)
    c61_62 = c61.merge(c62)
    c61_62.text = '點源(%)'

    c70 = table_2.cell(7, 0)
    c71 = table_2.cell(7, 1)
    c72 = table_2.cell(7, 2)
    c71_72 = c71.merge(c72)
    c71_72.text = '線源(%)'

    c80 = table_2.cell(8, 0)
    c81 = table_2.cell(8, 1)
    c82 = table_2.cell(8, 2)
    c81_82 = c81.merge(c82)
    c81_82.text = '面源(%)'
    c50_60_70_80 = c50.merge(c60).merge(c70).merge(c80)

    table_2.cell(6, 3).text = station.iloc[:, 51].astype(str) #P_SOX比例
    table_2.cell(6, 4).text = station.iloc[:, 52].astype(str) #P_NOX比例
    table_2.cell(6, 5).text = station.iloc[:, 53].astype(str) #P_THC比例
    table_2.cell(6, 6).text = station.iloc[:, 54].astype(str) #P_PM2.5比例

    table_2.cell(7, 3).text = station.iloc[:, 55].astype(str) #L_SOX比例
    table_2.cell(7, 4).text = station.iloc[:, 56].astype(str) #L_NOX比例
    table_2.cell(7, 5).text = station.iloc[:, 57].astype(str) #L_THC比例
    table_2.cell(7, 6).text = station.iloc[:, 58].astype(str) #L_PM2.5比例

    table_2.cell(8, 3).text = station.iloc[:, 59].astype(str) #A_SOX比例
    table_2.cell(8, 4).text = station.iloc[:, 60].astype(str) #A_NOX比例
    table_2.cell(8, 5).text = station.iloc[:, 61].astype(str) #A_THC比例
    table_2.cell(8, 6).text = station.iloc[:, 62].astype(str) #A_PM2.5比例

    c90 = table_2.cell(9, 0)
    c91 = table_2.cell(9, 1)
    c92 = table_2.cell(9, 2)
    c93 = table_2.cell(9, 3)
    c94 = table_2.cell(9, 4)
    c95 = table_2.cell(9, 5)
    c96 = table_2.cell(9, 6)
    c9 = c90.merge(c91).merge(c92).merge(c93).merge(c94).merge(c95).merge(c96)
    c9.text = '鄰近10公里最大排放量之廠商名稱、排放量、距離及方位'
    c9_xml_element = c9._tc #設定網底
    c9_properties = c9_xml_element.get_or_add_tcPr()
    shade_object = OxmlElement('w:shd')
    shade_object.set(qn('w:fill'), '#D9D9D9')
    c9_properties.append(shade_object)

    c10_0 = table_2.cell(10, 0)
    c10_1 = table_2.cell(10, 1)
    c10_2 = table_2.cell(10, 2)
    c10_3 = table_2.cell(10, 3)
    c10_4 = table_2.cell(10, 4)
    c10_5 = table_2.cell(10, 5)
    c10_6 = table_2.cell(10, 6)
    c10 = c10_0.merge(c10_1).merge(c10_2).merge(c10_3).merge(c10_4).merge(c10_5).merge(c10_6)
    c10.text = 'SOx'
    c10_xml_element = c10._tc #設定網底
    c10_properties = c10_xml_element.get_or_add_tcPr()
    shade_object = OxmlElement('w:shd')
    shade_object.set(qn('w:fill'), '#F2F2F2')
    c10_properties.append(shade_object)

    c11_0 = table_2.cell(11, 0)
    c11_1 = table_2.cell(11, 1)
    c11_2 = table_2.cell(11, 2)
    c11_3 = table_2.cell(11, 3)
    c11_4 = table_2.cell(11, 4)
    c11_5 = table_2.cell(11, 5)
    c11_6 = table_2.cell(11, 6)
    c110_111_112 = c11_0.merge(c11_1).merge(c11_2)
    c113_114 = c11_3.merge(c11_4)
    c110_111_112.text = '廠商名稱'
    c113_114.text = '排放量(T/y)'
    c11_5.text = '距離(km)'
    c11_6.text = '方位'

    c12_0 = table_2.cell(12, 0)
    c12_1 = table_2.cell(12, 1)
    c12_2 = table_2.cell(12, 2)
    c12_3 = table_2.cell(12, 3)
    c12_4 = table_2.cell(12, 4)
    c12_5 = table_2.cell(12, 5)
    c12_6 = table_2.cell(12, 6)
    c120_121_122 = c12_0.merge(c12_1).merge(c12_2)
    c123_124 = c12_3.merge(c12_4)

    c120_121_122.text = station.iloc[:, 63].astype(str) #SOX_fac
    c123_124.text = station.iloc[:, 64].astype(str) #SOX_emi
    c12_5.text = station.iloc[:, 66].astype(str) #distance
    c12_6.text = station.iloc[:, 65].astype(str) #position

    c13_0 = table_2.cell(13, 0)
    c13_1 = table_2.cell(13, 1)
    c13_2 = table_2.cell(13, 2)
    c13_3 = table_2.cell(13, 3)
    c13_4 = table_2.cell(13, 4)
    c13_5 = table_2.cell(13, 5)
    c13_6 = table_2.cell(13, 6)
    c13 = c13_0.merge(c13_1).merge(c13_2).merge(c13_3).merge(c13_4).merge(c13_5).merge(c13_6)
    c13.text = 'NOx'
    c13_xml_element = c13._tc #設定網底
    c13_properties = c13_xml_element.get_or_add_tcPr()
    shade_object = OxmlElement('w:shd')
    shade_object.set(qn('w:fill'), '#F2F2F2')
    c13_properties.append(shade_object)

    c14_0 = table_2.cell(14, 0)
    c14_1 = table_2.cell(14, 1)
    c14_2 = table_2.cell(14, 2)
    c14_3 = table_2.cell(14, 3)
    c14_4 = table_2.cell(14, 4)
    c14_5 = table_2.cell(14, 5)
    c14_6 = table_2.cell(14, 6)
    c140_141_142 = c14_0.merge(c14_1).merge(c14_2)
    c143_144 = c14_3.merge(c14_4)
    c140_141_142.text = '廠商名稱'
    c143_144.text = '排放量(T/y)'
    c14_5.text = '距離(km)'
    c14_6.text = '方位'

    c15_0 = table_2.cell(15, 0)
    c15_1 = table_2.cell(15, 1)
    c15_2 = table_2.cell(15, 2)
    c15_3 = table_2.cell(15, 3)
    c15_4 = table_2.cell(15, 4)
    c15_5 = table_2.cell(15, 5)
    c15_6 = table_2.cell(15, 6)
    c150_151_152 = c15_0.merge(c15_1).merge(c15_2)
    c153_154 = c15_3.merge(c15_4)

    c150_151_152.text = station.iloc[:, 67].astype(str) #NOX_fac
    c153_154.text = station.iloc[:, 68].astype(str) #NOX_emi
    c15_5.text = station.iloc[:, 70].astype(str) #distance
    c15_6.text = station.iloc[:, 69].astype(str) #position

    c160 = table_2.cell(16, 0)
    c161 = table_2.cell(16, 1)
    c162 = table_2.cell(16, 2)
    c163 = table_2.cell(16, 3)
    c164 = table_2.cell(16, 4)
    c165 = table_2.cell(16, 5)
    c166 = table_2.cell(16, 6)
    c16 = c160.merge(c161).merge(c162).merge(c163).merge(c164).merge(c165).merge(c166)
    c16.text = 'THC'
    c16_xml_element = c16._tc #設定網底
    c16_properties = c16_xml_element.get_or_add_tcPr()
    shade_object = OxmlElement('w:shd')
    shade_object.set(qn('w:fill'), '#F2F2F2')
    c16_properties.append(shade_object)

    c17_0 = table_2.cell(17, 0)
    c17_1 = table_2.cell(17, 1)
    c17_2 = table_2.cell(17, 2)
    c17_3 = table_2.cell(17, 3)
    c17_4 = table_2.cell(17, 4)
    c17_5 = table_2.cell(17, 5)
    c17_6 = table_2.cell(17, 6)
    c170_171_172 = c17_0.merge(c17_1).merge(c17_2)
    c173_174 = c17_3.merge(c17_4)
    c170_171_172.text = '廠商名稱'
    c173_174.text = '排放量(T/y)'
    c17_5.text = '距離(km)'
    c17_6.text = '方位'

    c18_0 = table_2.cell(18, 0)
    c18_1 = table_2.cell(18, 1)
    c18_2 = table_2.cell(18, 2)
    c18_3 = table_2.cell(18, 3)
    c18_4 = table_2.cell(18, 4)
    c18_5 = table_2.cell(18, 5)
    c18_6 = table_2.cell(18, 6)
    c180_181_182 = c18_0.merge(c18_1).merge(c18_2)
    c183_184 = c18_3.merge(c18_4)

    c180_181_182.text = station.iloc[:, 71].astype(str) #THC_fac
    c183_184.text = station.iloc[:, 72].astype(str) #THC_emi
    c18_5.text = station.iloc[:, 74].astype(str) #distance
    c18_6.text = station.iloc[:, 73].astype(str) #position

    c190 = table_2.cell(19, 0)
    c191 = table_2.cell(19, 1)
    c192 = table_2.cell(19, 2)
    c193 = table_2.cell(19, 3)
    c194 = table_2.cell(19, 4)
    c195 = table_2.cell(19, 5)
    c196 = table_2.cell(19, 6)
    c19 = c190.merge(c191).merge(c192).merge(c193).merge(c194).merge(c195).merge(c196)
    c19.text = 'TSP'
    c19_xml_element = c19._tc #設定網底
    c19_properties = c19_xml_element.get_or_add_tcPr()
    shade_object = OxmlElement('w:shd')
    shade_object.set(qn('w:fill'), '#F2F2F2')
    c19_properties.append(shade_object)

    c20_0 = table_2.cell(20, 0)
    c20_1 = table_2.cell(20, 1)
    c20_2 = table_2.cell(20, 2)
    c20_3 = table_2.cell(20, 3)
    c20_4 = table_2.cell(20, 4)
    c20_5 = table_2.cell(20, 5)
    c20_6 = table_2.cell(20, 6)
    c200_201_202 = c20_0.merge(c20_1).merge(c20_2)
    c203_204 = c20_3.merge(c20_4)
    c200_201_202.text = '廠商名稱'
    c203_204.text = '排放量(T/y)'
    c20_5.text = '距離(km)'
    c20_6.text = '方位'

    c21_0 = table_2.cell(21, 0)
    c21_1 = table_2.cell(21, 1)
    c21_2 = table_2.cell(21, 2)
    c21_3 = table_2.cell(21, 3)
    c21_4 = table_2.cell(21, 4)
    c21_5 = table_2.cell(21, 5)
    c21_6 = table_2.cell(21, 6)
    c210_211_212 = c21_0.merge(c21_1).merge(c21_2)
    c213_214 = c21_3.merge(c21_4)

    c210_211_212.text = station.iloc[:, 75].astype(str) #TSP_fac
    c213_214.text = station.iloc[:, 76].astype(str) #TSP_emi
    c21_5.text = station.iloc[:, 78].astype(str) #distance
    c21_6.text = station.iloc[:, 77].astype(str) #position


    document.add_page_break()
    
    table_2.cell(21, 0).width = Inches(3)
    table_2.cell(21, 1).width = Inches(2)
    table_2.cell(21, 2).width = Inches(1)
    table_2.cell(21, 3).width = Inches(1)
    table_2.cell(21, 4).width = Inches(1)
    table_2.cell(21, 5).width = Inches(1)
    table_2.cell(21, 6).width = Inches(1)
#########################表格邊框設定#############################
    table_1.style = 'Table Grid'
    table_2.style = 'Table Grid'
#########################表格高度設定#############################
    for row in table_1.rows:
        row.height = Cm(0.6)

    for row in table_2.rows:
        row.height = Cm(0.6)
#########################修改表格內字體大小、字型及取代內容等#############################
    for row in table_1.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #表格內容水平置中
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #表格內容垂直置中
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(9)
                    font.bold = True
                    run.font.name = '微軟正黑體'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微軟正黑體') # 中文字型要多加這一列
                    if 'nan' in run.text: #將表格內有nan的部分以空白取代
                        run.text = run.text.replace('nan', '-')

    for row in table_2.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER #表格內容水平置中
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #表格內容垂直置中
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(9)
                    font.bold = True
                    run.font.name = '微軟正黑體'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微軟正黑體') # 中文字型要多加這一列
                    if 'nan' in run.text: #將表格內有nan的部分以空白取代
                        run.text = run.text.replace('nan', '-')
#########################儲存#############################
    document.save('環保署測站基本資料表_年報用.docx') #儲存